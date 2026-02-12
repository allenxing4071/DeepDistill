"""
Google Drive 多格式导出器
支持三种导出格式：
  - doc: Markdown -> HTML -> Google Doc（在线编辑）
  - word: Markdown -> .docx（python-docx 生成）-> 上传到 Google Drive
  - excel: 结构化数据 -> .xlsx（openpyxl 生成）-> 上传到 Google Drive
认证方式：OAuth2（首次使用需浏览器授权，token 自动缓存刷新）。
"""

from __future__ import annotations

import io
import json
import logging
import tempfile
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

import markdown

logger = logging.getLogger("deepdistill.export.google_docs")

# Google Drive API 上传重试配置
GDRIVE_MAX_RETRIES = 3
GDRIVE_RETRY_DELAY = 2  # 秒，指数退避基数

# Google API 所需权限范围
SCOPES = [
    "https://www.googleapis.com/auth/drive.file",  # 仅管理本应用创建的文件
]


class GoogleDocsExporter:
    """Google Docs 导出器：Markdown -> HTML -> Google Doc，支持按分类子文件夹管理"""

    # 预定义分类（中文名 -> 英文文件夹名）
    CATEGORIES = {
        "投诉维权": "投诉维权",
        "学习笔记": "学习笔记",
        "技术文档": "技术文档",
        "市场分析": "市场分析",
        "会议纪要": "会议纪要",
        "创意素材": "创意素材",
        "法律法规": "法律法规",
        "其他": "其他",
    }

    def __init__(
        self,
        credentials_path: str | Path,
        token_path: str | Path,
        folder_name: str = "DeepDistill",
    ):
        self.credentials_path = Path(credentials_path)
        self.token_path = Path(token_path)
        self.folder_name = folder_name
        self._drive_service = None
        self._folder_id: str | None = None
        self._subfolder_cache: dict[str, str] = {}  # category -> folder_id

    def _authenticate(self):
        """OAuth2 认证：加载缓存 token 或启动浏览器授权流程"""
        from google.auth.transport.requests import Request
        from google.oauth2.credentials import Credentials
        from google_auth_oauthlib.flow import InstalledAppFlow

        creds = None

        # 尝试加载缓存的 token
        if self.token_path.exists():
            try:
                creds = Credentials.from_authorized_user_file(str(self.token_path), SCOPES)
            except Exception as e:
                logger.warning(f"加载缓存 token 失败: {e}，将重新授权")
                creds = None

        # token 过期则刷新
        if creds and creds.expired and creds.refresh_token:
            try:
                creds.refresh(Request())
                logger.info("Google OAuth2 token 已刷新")
            except Exception as e:
                logger.warning(f"刷新 token 失败: {e}，将重新授权")
                creds = None

        # 无有效 token，启动浏览器授权
        if not creds or not creds.valid:
            if not self.credentials_path.exists():
                raise FileNotFoundError(
                    f"Google OAuth2 凭据文件不存在: {self.credentials_path}\n"
                    "请从 Google Cloud Console 下载 OAuth2 Client ID JSON 文件，"
                    "放到 config/google_credentials.json"
                )
            flow = InstalledAppFlow.from_client_secrets_file(
                str(self.credentials_path), SCOPES
            )
            # 在服务器环境中使用 port=0 自动选择端口
            creds = flow.run_local_server(port=0, open_browser=True)
            logger.info("Google OAuth2 授权成功")

        # 缓存 token
        self.token_path.parent.mkdir(parents=True, exist_ok=True)
        with open(self.token_path, "w") as f:
            f.write(creds.to_json())

        return creds

    def _get_drive_service(self):
        """获取 Google Drive API 服务实例（懒加载 + 缓存）"""
        if self._drive_service is None:
            from googleapiclient.discovery import build

            creds = self._authenticate()
            self._drive_service = build("drive", "v3", credentials=creds)
            logger.info("Google Drive API 服务已初始化")
        return self._drive_service

    def _ensure_folder(self) -> str:
        """确保 Google Drive 中存在目标文件夹，返回文件夹 ID"""
        if self._folder_id:
            return self._folder_id

        service = self._get_drive_service()

        # 搜索已有文件夹
        query = (
            f"name = '{self.folder_name}' "
            "and mimeType = 'application/vnd.google-apps.folder' "
            "and trashed = false"
        )
        results = service.files().list(
            q=query, spaces="drive", fields="files(id, name)", pageSize=1
        ).execute()

        files = results.get("files", [])
        if files:
            self._folder_id = files[0]["id"]
            logger.info(f"找到已有文件夹 '{self.folder_name}' (ID: {self._folder_id})")
            return self._folder_id

        # 创建新文件夹
        folder_metadata = {
            "name": self.folder_name,
            "mimeType": "application/vnd.google-apps.folder",
        }
        folder = service.files().create(
            body=folder_metadata, fields="id"
        ).execute()
        self._folder_id = folder["id"]
        logger.info(f"创建文件夹 '{self.folder_name}' (ID: {self._folder_id})")
        return self._folder_id

    def _ensure_subfolder(self, category: str) -> str:
        """确保分类子文件夹存在，返回子文件夹 ID"""
        if category in self._subfolder_cache:
            return self._subfolder_cache[category]

        root_id = self._ensure_folder()
        service = self._get_drive_service()

        # 搜索已有子文件夹
        query = (
            f"name = '{category}' "
            f"and '{root_id}' in parents "
            "and mimeType = 'application/vnd.google-apps.folder' "
            "and trashed = false"
        )
        results = service.files().list(
            q=query, spaces="drive", fields="files(id, name)", pageSize=1
        ).execute()

        files = results.get("files", [])
        if files:
            folder_id = files[0]["id"]
            logger.info(f"找到已有子文件夹 '{category}' (ID: {folder_id})")
        else:
            # 创建子文件夹
            folder_metadata = {
                "name": category,
                "mimeType": "application/vnd.google-apps.folder",
                "parents": [root_id],
            }
            folder = service.files().create(
                body=folder_metadata, fields="id"
            ).execute()
            folder_id = folder["id"]
            logger.info(f"创建子文件夹 '{category}' (ID: {folder_id})")

        self._subfolder_cache[category] = folder_id
        return folder_id

    def list_categories(self) -> list[dict]:
        """列出所有预定义分类及其在 Google Drive 中的文档数量"""
        root_id = self._ensure_folder()
        service = self._get_drive_service()
        result = []

        for cat_name in self.CATEGORIES:
            # 查找子文件夹
            query = (
                f"name = '{cat_name}' "
                f"and '{root_id}' in parents "
                "and mimeType = 'application/vnd.google-apps.folder' "
                "and trashed = false"
            )
            folders = service.files().list(
                q=query, spaces="drive", fields="files(id)", pageSize=1
            ).execute().get("files", [])

            doc_count = 0
            folder_url = None
            if folders:
                fid = folders[0]["id"]
                # 统计该文件夹下的文档数
                doc_query = (
                    f"'{fid}' in parents "
                    "and mimeType = 'application/vnd.google-apps.document' "
                    "and trashed = false"
                )
                docs = service.files().list(
                    q=doc_query, spaces="drive", fields="files(id)", pageSize=100
                ).execute().get("files", [])
                doc_count = len(docs)
                folder_url = f"https://drive.google.com/drive/folders/{fid}"

            result.append({
                "name": cat_name,
                "doc_count": doc_count,
                "folder_url": folder_url,
            })

        return result

    def _markdown_to_html(self, md_content: str) -> str:
        """将 Markdown 转为带基本样式的 HTML"""
        html_body = markdown.markdown(
            md_content,
            extensions=["tables", "fenced_code", "toc"],
        )
        # 包装为完整 HTML 文档，带基本样式
        html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
body {{ font-family: 'Noto Sans SC', Arial, sans-serif; line-height: 1.6; color: #333; }}
h1 {{ color: #1a1a1a; border-bottom: 2px solid #4285f4; padding-bottom: 8px; }}
h2 {{ color: #333; margin-top: 24px; }}
h3 {{ color: #555; }}
code {{ background: #f5f5f5; padding: 2px 6px; border-radius: 3px; font-size: 0.9em; }}
pre {{ background: #f5f5f5; padding: 12px; border-radius: 6px; overflow-x: auto; }}
blockquote {{ border-left: 3px solid #4285f4; padding-left: 12px; color: #666; margin: 12px 0; }}
table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}
th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
th {{ background: #f5f5f5; font-weight: bold; }}
</style>
</head>
<body>
{html_body}
</body>
</html>"""
        return html

    def export_markdown(self, md_content: str, title: str, category: str | None = None) -> dict:
        """
        将 Markdown 内容导出为 Google Doc。

        Args:
            md_content: Markdown 格式的蒸馏结果
            title: 文档标题
            category: 分类名称（可选），指定后文档放入对应子文件夹

        Returns:
            {"doc_id": str, "doc_url": str, "title": str, "category": str | None}
        """
        from googleapiclient.http import MediaInMemoryUpload

        service = self._get_drive_service()

        # 强制放入子目录，不允许根目录
        if not category or category not in self.CATEGORIES:
            category = "其他"
            logger.info(f"未指定有效分类，默认放入「其他」目录")
        folder_id = self._ensure_subfolder(category)

        # Markdown -> HTML
        html_content = self._markdown_to_html(md_content)

        # 上传 HTML 并自动转为 Google Doc
        file_metadata = {
            "name": title,
            "mimeType": "application/vnd.google-apps.document",
            "parents": [folder_id],
        }
        media = MediaInMemoryUpload(
            html_content.encode("utf-8"),
            mimetype="text/html",
            resumable=True,
        )

        # 带重试的上传（防止网络抖动/API 限流）
        last_error = None
        for attempt in range(GDRIVE_MAX_RETRIES):
            try:
                file = service.files().create(
                    body=file_metadata,
                    media_body=media,
                    fields="id, webViewLink, name",
                ).execute()
                break
            except Exception as e:
                last_error = e
                if attempt < GDRIVE_MAX_RETRIES - 1:
                    wait = GDRIVE_RETRY_DELAY * (2 ** attempt)
                    logger.warning(f"Google Drive 上传失败（第 {attempt + 1} 次）: {e}，{wait}s 后重试")
                    time.sleep(wait)
                else:
                    raise RuntimeError(f"Google Drive 上传失败（已重试 {GDRIVE_MAX_RETRIES} 次）: {last_error}")

        doc_id = file["id"]
        doc_url = file.get("webViewLink", f"https://docs.google.com/document/d/{doc_id}/edit")

        logger.info(f"已导出到 Google Docs: {title} -> {doc_url}")

        return {
            "doc_id": doc_id,
            "doc_url": doc_url,
            "title": title,
            "category": category,
            "folder_url": f"https://drive.google.com/drive/folders/{folder_id}",
        }

    def export_from_file(self, md_file_path: str | Path, category: str | None = None) -> dict:
        """从 Markdown 文件导出到 Google Doc"""
        md_path = Path(md_file_path)
        if not md_path.exists():
            raise FileNotFoundError(f"Markdown 文件不存在: {md_path}")

        md_content = md_path.read_text(encoding="utf-8")
        title = md_path.stem.replace("_distilled", "").replace("_skill", "")
        return self.export_markdown(md_content, title, category=category)

    @staticmethod
    def _generate_short_title(task: dict) -> str:
        """
        从 AI 分析结果中提取 ≤8 字的中文短标题。
        优先级：summary 中文摘要 > key_points 第一条 > 中文 keywords > 文件名
        强制中文输出：如果提取到英文，则从 summary 中截取中文部分。
        """
        import re

        result = task.get("result", {})
        ai = result.get("ai_result") or result.get("ai_analysis") or {}
        filename = task.get("filename", "未知文件")

        def _has_chinese(text: str) -> bool:
            """检测文本是否包含中文字符"""
            return bool(re.search(r'[\u4e00-\u9fff]', text))

        def _extract_from_summary(summary: str) -> str | None:
            """从 summary 中提取 ≤8 字的中文短标题"""
            if not summary:
                return None
            # 去掉开头的套话
            summary = re.sub(
                r'^(本文|该文|这篇文章|文章|本视频|该视频|这个视频|视频|本页面|该页面|页面)'
                r'(主要|详细|全面|系统|深入)?'
                r'(介绍|讲解|分析|阐述|探讨|说明|描述|总结|概述|讨论|涵盖|涉及|关注|聚焦)'
                r'(了)?',
                '', summary
            )
            # 跳过开头的英文/数字/空格/标点，找到第一个中文字符
            cn_start = re.search(r'[\u4e00-\u9fff]', summary)
            if cn_start:
                summary = summary[cn_start.start():]
            # 跳过开头的虚词（"的/在/了/是/有/和/与"等）
            summary = re.sub(r'^[的在了是有和与及]', '', summary)
            # 去除夹杂的英文单词（保留中文连续片段）
            # 提取所有中文连续片段
            cn_segments = re.findall(r'[\u4e00-\u9fff]+', summary[:30])
            if cn_segments:
                # 拼接前几个中文片段直到 ≤8 字
                title = ""
                for seg in cn_segments:
                    if len(title) + len(seg) <= 8:
                        title += seg
                    else:
                        # 截取部分
                        remaining = 8 - len(title)
                        if remaining >= 2:
                            title += seg[:remaining]
                        break
                # 在虚词处优化截断
                if len(title) > 6:
                    for cut_word in ["和", "与", "的", "及", "等", "在", "方面"]:
                        idx = title.rfind(cut_word, 3)
                        if idx >= 3:
                            title = title[:idx]
                            break
                if title and len(title) >= 2:
                    return title
            return None

        # 1. 优先从 summary（中文摘要）提取
        summary = ai.get("summary", "")
        title = _extract_from_summary(summary)
        if title:
            return title

        # 2. 尝试从 key_points 第一条提取
        key_points = ai.get("key_points", [])
        if key_points:
            first_point = key_points[0]
            kp_title = _extract_from_summary(first_point)
            if kp_title:
                return kp_title

        # 3. 尝试从 keywords 中找中文关键词
        keywords = ai.get("keywords", [])
        for kw in keywords:
            kw = kw.strip().strip("#").strip("`")
            if kw and _has_chinese(kw) and len(kw) <= 8:
                return kw

        # 4. 如果 summary 有内容但全是英文，提取核心名词短语
        if summary:
            # 尝试找到第一个有意义的中文片段
            cn_match = re.search(r'[\u4e00-\u9fff]{2,8}', summary)
            if cn_match:
                return cn_match.group()

        # 5. fallback：从文件名提取
        stem = Path(filename).stem
        if _has_chinese(stem):
            return stem[:8] if len(stem) > 8 else stem

        # 6. 最终 fallback：使用"未命名文档"
        return "未命名文档"

    def _build_doc_markdown(self, task: dict) -> tuple[str, str]:
        """构建普通文档格式的 Markdown，返回 (md_content, title)"""
        result = task.get("result", {})
        filename = task.get("filename", "未知文件")
        title = self._generate_short_title(task)

        # 优先使用已生成的 Markdown 文件
        output_path = result.get("output_path", "")
        if output_path and Path(output_path).exists() and output_path.endswith(".md"):
            return Path(output_path).read_text(encoding="utf-8"), title

        # 从 result 数据重新生成
        md_lines = [f"# {title}", ""]
        ai_result = result.get("ai_result") or result.get("ai_analysis") or {}

        if ai_result.get("summary"):
            md_lines += ["## 摘要", "", ai_result["summary"], ""]
        if ai_result.get("key_points"):
            md_lines += ["## 核心观点", ""]
            md_lines += [f"- {p}" for p in ai_result["key_points"]]
            md_lines.append("")
        if ai_result.get("keywords"):
            md_lines += ["## 关键词", "", " ".join(f"`{kw}`" for kw in ai_result["keywords"]), ""]

        # 普通文档中不再包含原始文本（已有独立的 [源文件] 文档）
        raw_text = result.get("extracted_text") or result.get("raw_text", "")
        if raw_text:
            preview = raw_text[:500] + ("..." if len(raw_text) > 500 else "")
            md_lines += ["---", "", "## 原始文本预览", "", f"> 完整原始文本请查看同目录下的 **[源文件]** 文档（共 {len(raw_text)} 字符）", "", preview, ""]

        return "\n".join(md_lines), stem

    def _build_skill_markdown(self, task: dict) -> tuple[str, str]:
        """构建 Skill 文档格式的 Markdown，返回 (md_content, title)"""
        result = task.get("result", {})
        filename = task.get("filename", "未知文件")
        title = self._generate_short_title(task)
        ai = result.get("ai_result") or result.get("ai_analysis") or {}
        now = datetime.now(timezone.utc).strftime("%Y-%m-%d")

        lines = []
        # 标题
        lines += [f"# SKILL: {title}", ""]

        # 核心摘要
        summary = ai.get("summary", "")
        if summary:
            lines += [f"> **核心摘要**: {summary}", ""]

        # 关键词标签
        keywords = ai.get("keywords", [])
        if keywords:
            lines += ["**标签**: " + " ".join(f"`#{kw}`" for kw in keywords), ""]

        # 知识要点
        key_points = ai.get("key_points", [])
        if key_points:
            lines += ["## 知识要点", ""]
            for i, point in enumerate(key_points, 1):
                lines.append(f"{i}. {point}")
            lines.append("")

        # 内容结构
        structure = ai.get("structure", {})
        sections = structure.get("sections", [])
        if sections:
            lines += ["## 详细内容", ""]
            for section in sections:
                heading = section.get("heading", "未命名")
                content = section.get("content", "")
                lines += [f"### {heading}", "", content, ""]

        # 实践指南
        lines += ["## 实践指南", ""]
        lines += ["1. 快速浏览核心摘要了解主旨"]
        lines += ["2. 根据关键词标签关联相关知识"]
        lines += ["3. 深入阅读感兴趣的章节"]
        lines += ["4. 将知识要点应用到实际项目中"]
        lines.append("")

        # 关联知识
        if keywords:
            lines += ["## 关联知识", ""]
            for kw in keywords[:5]:
                lines.append(f"- 搜索: `{kw}` 查找相关内容")
            lines.append("")

        # 元信息
        source_type = result.get("source_type", "未知")
        text_len = result.get("extracted_text_length", 0)
        lines += ["---", "", "## 元信息", ""]
        lines += [f"- **来源文件**: `{filename}`"]
        lines += [f"- **文件类型**: {source_type}"]
        lines += [f"- **提取文本长度**: {text_len} 字符"]
        lines += [f"- **生成时间**: {now}"]
        lines.append("")

        return "\n".join(lines), f"{title} [SKILL]"

    def _build_raw_markdown(self, task: dict) -> tuple[str, str]:
        """构建原始源文件的 Markdown（完整文本，不截断），返回 (md_content, title)"""
        result = task.get("result", {})
        filename = task.get("filename", "未知文件")
        title = self._generate_short_title(task)
        now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M")

        raw_text = result.get("extracted_text") or result.get("raw_text", "")

        md_lines = [
            f"# {title} — 源文件",
            "",
            "> ⚠️ **这是原始提取文本（未经 AI 加工），完整保留了源内容。**",
            "",
            "---",
            "",
            f"- **来源文件**: `{filename}`",
            f"- **文件类型**: {result.get('source_type', '未知')}",
            f"- **文本长度**: {len(raw_text)} 字符",
            f"- **提取时间**: {now}",
            "",
            "---",
            "",
            "## 原始文本",
            "",
            raw_text if raw_text else "（无提取文本）",
            "",
        ]

        return "\n".join(md_lines), f"{title} [源文件]"

    @staticmethod
    def _auto_categorize(task: dict) -> str:
        """
        根据 AI 分析结果自动推断分类。
        关键词匹配 → 默认"其他"。保证永远不返回 None。
        """
        result = task.get("result", {})
        ai = result.get("ai_result") or result.get("ai_analysis") or {}
        keywords = [kw.lower() for kw in (ai.get("keywords") or [])]
        summary = (ai.get("summary") or "").lower()
        text_blob = " ".join(keywords) + " " + summary

        # 关键词 → 分类映射
        category_keywords = {
            "技术文档": ["api", "code", "docker", "python", "javascript", "react", "框架", "编程",
                       "算法", "架构", "部署", "开发", "技术", "软件", "数据库", "服务器",
                       "kubernetes", "linux", "git", "machine learning", "deep learning",
                       "ai", "artificial intelligence", "nlp", "computer vision"],
            "市场分析": ["市场", "行情", "交易", "投资", "金融", "股票", "加密", "比特币",
                       "以太坊", "区块链", "crypto", "bitcoin", "ethereum", "blockchain",
                       "cryptocurrency", "trading", "finance", "defi"],
            "学习笔记": ["教程", "学习", "入门", "指南", "tutorial", "guide", "course",
                       "documentation", "docs", "笔记", "总结", "知识"],
            "投诉维权": ["投诉", "维权", "举报", "违规", "欺诈", "诈骗"],
            "会议纪要": ["会议", "纪要", "讨论", "决议", "meeting", "minutes"],
            "创意素材": ["设计", "素材", "图片", "视频", "创意", "风格", "配色", "ui", "ux"],
            "法律法规": ["法律", "法规", "条例", "合规", "监管", "regulation", "law", "legal"],
        }

        for cat, kws in category_keywords.items():
            if any(kw in text_blob for kw in kws):
                return cat

        return "其他"

    def export_task_result(
        self, task: dict, category: str | None = None,
        fmt: str = "doc", export_format: str = "doc",
    ) -> dict | list[dict]:
        """
        从任务结果导出到 Google Drive。
        每次导出都会额外生成一份 [源文件]（完整原始文本，不截断）。
        文件名使用 AI 提炼的 ≤8 字中文标题。
        强制放入分类子目录（不允许根目录）。

        Args:
            task: 任务数据
            category: 分类子文件夹名称（为 None 时自动推断）
            fmt: 文档类型 — "doc"(普通文档) / "skill"(Skill文档) / "both"(两者都导出)
            export_format: 文件格式 — "doc"(Google Doc) / "word"(.docx) / "excel"(.xlsx)

        Returns:
            始终返回 list[dict]，最后一项为 [源文件]
        """
        result = task.get("result")
        if not result:
            raise ValueError("任务尚未完成或无结果")

        # 强制分类：如果未指定或不在预定义列表中，自动推断
        if not category or category not in self.CATEGORIES:
            category = self._auto_categorize(task)
            logger.info(f"自动分类: {category}")

        def _export_one(md_content: str, title: str) -> dict:
            """根据 export_format 选择导出方式"""
            if export_format == "word":
                return self._export_as_word(md_content, title, category=category)
            elif export_format == "excel":
                return self._export_as_excel(task, title, category=category)
            else:
                return self.export_markdown(md_content, title, category=category)

        results = []

        if fmt == "both":
            md_doc, title_doc = self._build_doc_markdown(task)
            results.append(_export_one(md_doc, title_doc))

            md_skill, title_skill = self._build_skill_markdown(task)
            results.append(_export_one(md_skill, title_skill))

        elif fmt == "skill":
            md_content, title = self._build_skill_markdown(task)
            results.append(_export_one(md_content, title))
        else:
            md_content, title = self._build_doc_markdown(task)
            results.append(_export_one(md_content, title))

        # ── 额外导出源文件（完整原始文本，始终用 Google Doc 格式） ──
        raw_text = (result.get("extracted_text") or result.get("raw_text", ""))
        if raw_text:
            md_raw, title_raw = self._build_raw_markdown(task)
            raw_result = self.export_markdown(md_raw, title_raw, category=category)
            raw_result["is_raw"] = True  # 标记为源文件
            results.append(raw_result)
            logger.info(f"已导出源文件: {title_raw}")

        return results

    # ── Word (.docx) 导出 ──

    def _export_as_word(
        self, md_content: str, title: str, category: str | None = None,
    ) -> dict:
        """
        将 Markdown 内容生成 .docx 文件并上传到 Google Drive。
        使用 python-docx 库生成 Word 文档。
        """
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(11)

        # 解析 Markdown 并写入 Word 文档
        lines = md_content.split("\n")
        for line in lines:
            stripped = line.strip()

            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("- "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped.startswith("1. ") or stripped.startswith("2. ") or stripped.startswith("3. "):
                # 有序列表
                text = stripped.split(". ", 1)[1] if ". " in stripped else stripped
                doc.add_paragraph(text, style="List Number")
            elif stripped.startswith("> "):
                # 引用块 — 使用斜体段落
                p = doc.add_paragraph()
                run = p.add_run(stripped[2:])
                run.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)
            elif stripped == "---":
                # 分隔线 — 添加空段落
                doc.add_paragraph()
            elif stripped:
                doc.add_paragraph(stripped)
            # 空行跳过

        # 保存到内存缓冲区
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # 上传到 Google Drive
        return self._upload_binary_to_drive(
            buffer.read(),
            title=f"{title}.docx",
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            category=category,
        )

    # ── Excel (.xlsx) 导出 ──

    def _export_as_excel(
        self, task: dict, title: str, category: str | None = None,
    ) -> dict:
        """
        将任务结构化数据生成 .xlsx 文件并上传到 Google Drive。
        使用 openpyxl 库生成 Excel 文档。
        """
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        result = task.get("result", {})
        ai = result.get("ai_result") or result.get("ai_analysis") or {}
        filename = task.get("filename", "未知文件")

        wb = Workbook()

        # ── Sheet 1: 摘要 ──
        ws_summary = wb.active
        ws_summary.title = "摘要"

        # 样式定义
        header_font = Font(bold=True, size=12, color="FFFFFF")
        header_fill = PatternFill(start_color="4285F4", end_color="4285F4", fill_type="solid")
        thin_border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )

        # 标题行
        ws_summary.append(["字段", "内容"])
        for cell in ws_summary[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.border = thin_border
            cell.alignment = Alignment(horizontal="center")

        # 数据行
        rows = [
            ("文件名", filename),
            ("文件类型", result.get("source_type", "未知")),
            ("提取文本长度", str(result.get("extracted_text_length", 0)) + " 字符"),
            ("摘要", ai.get("summary", "")),
        ]
        for field, value in rows:
            ws_summary.append([field, value])
            for cell in ws_summary[ws_summary.max_row]:
                cell.border = thin_border
                cell.alignment = Alignment(wrap_text=True, vertical="top")

        # 列宽
        ws_summary.column_dimensions["A"].width = 18
        ws_summary.column_dimensions["B"].width = 80

        # ── Sheet 2: 核心观点 ──
        key_points = ai.get("key_points", [])
        if key_points:
            ws_points = wb.create_sheet("核心观点")
            ws_points.append(["序号", "观点"])
            for cell in ws_points[1]:
                cell.font = header_font
                cell.fill = header_fill
                cell.border = thin_border
            for i, point in enumerate(key_points, 1):
                ws_points.append([i, point])
                for cell in ws_points[ws_points.max_row]:
                    cell.border = thin_border
                    cell.alignment = Alignment(wrap_text=True, vertical="top")
            ws_points.column_dimensions["A"].width = 8
            ws_points.column_dimensions["B"].width = 80

        # ── Sheet 3: 关键词 ──
        keywords = ai.get("keywords", [])
        if keywords:
            ws_kw = wb.create_sheet("关键词")
            ws_kw.append(["序号", "关键词"])
            for cell in ws_kw[1]:
                cell.font = header_font
                cell.fill = header_fill
                cell.border = thin_border
            for i, kw in enumerate(keywords, 1):
                ws_kw.append([i, kw])
                for cell in ws_kw[ws_kw.max_row]:
                    cell.border = thin_border
            ws_kw.column_dimensions["A"].width = 8
            ws_kw.column_dimensions["B"].width = 30

        # ── Sheet 4: 内容结构（如有） ──
        structure = ai.get("structure", {})
        sections = structure.get("sections", [])
        if sections:
            ws_struct = wb.create_sheet("内容结构")
            ws_struct.append(["章节", "内容"])
            for cell in ws_struct[1]:
                cell.font = header_font
                cell.fill = header_fill
                cell.border = thin_border
            for section in sections:
                heading = section.get("heading", "")
                content = section.get("content", "")
                ws_struct.append([heading, content])
                for cell in ws_struct[ws_struct.max_row]:
                    cell.border = thin_border
                    cell.alignment = Alignment(wrap_text=True, vertical="top")
            ws_struct.column_dimensions["A"].width = 25
            ws_struct.column_dimensions["B"].width = 80

        # 保存到内存缓冲区
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        # 上传到 Google Drive
        return self._upload_binary_to_drive(
            buffer.read(),
            title=f"{title}.xlsx",
            mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            category=category,
        )

    # ── 通用二进制文件上传到 Google Drive ──

    def _upload_binary_to_drive(
        self, data: bytes, title: str, mime_type: str, category: str | None = None,
    ) -> dict:
        """
        将二进制文件上传到 Google Drive（不转换格式）。

        Args:
            data: 文件二进制内容
            title: 文件名（含扩展名）
            mime_type: MIME 类型
            category: 分类子文件夹名称

        Returns:
            {"doc_id": str, "doc_url": str, "title": str, "category": str | None}
        """
        from googleapiclient.http import MediaInMemoryUpload

        service = self._get_drive_service()

        # 强制放入子目录，不允许根目录
        if not category or category not in self.CATEGORIES:
            category = "其他"
            logger.info(f"未指定有效分类，默认放入「其他」目录")
        folder_id = self._ensure_subfolder(category)

        file_metadata = {
            "name": title,
            "parents": [folder_id],
        }
        media = MediaInMemoryUpload(data, mimetype=mime_type, resumable=True)

        # 带重试的上传
        last_error = None
        for attempt in range(GDRIVE_MAX_RETRIES):
            try:
                file = service.files().create(
                    body=file_metadata,
                    media_body=media,
                    fields="id, webViewLink, name",
                ).execute()
                break
            except Exception as e:
                last_error = e
                if attempt < GDRIVE_MAX_RETRIES - 1:
                    wait = GDRIVE_RETRY_DELAY * (2 ** attempt)
                    logger.warning(f"Google Drive 上传失败（第 {attempt + 1} 次）: {e}，{wait}s 后重试")
                    time.sleep(wait)
                else:
                    raise RuntimeError(f"Google Drive 上传失败（已重试 {GDRIVE_MAX_RETRIES} 次）: {last_error}")

        doc_id = file["id"]
        doc_url = file.get("webViewLink", f"https://drive.google.com/file/d/{doc_id}/view")

        logger.info(f"已上传到 Google Drive: {title} -> {doc_url}")

        return {
            "doc_id": doc_id,
            "doc_url": doc_url,
            "title": title,
            "category": category,
            "folder_url": f"https://drive.google.com/drive/folders/{folder_id}",
        }


def get_exporter() -> GoogleDocsExporter:
    """获取全局 Google Docs 导出器实例"""
    from ..config import cfg

    return GoogleDocsExporter(
        credentials_path=cfg.GOOGLE_DOCS_CREDENTIALS_PATH,
        token_path=cfg.GOOGLE_DOCS_TOKEN_PATH,
        folder_name=cfg.GOOGLE_DOCS_FOLDER_NAME,
    )
